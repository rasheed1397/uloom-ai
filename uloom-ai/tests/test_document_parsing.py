from io import BytesIO

import docx
import pytest
from pypdf import PdfWriter

from app.services.document_parsing import UnsupportedMimeTypeError, extract_text


def test_extract_text_plain():
    segments = extract_text("text/plain", b"Hello world.")

    assert len(segments) == 1
    assert segments[0].text == "Hello world."
    assert segments[0].source_location == {"offset": 0}


def test_extract_text_markdown_uses_plain_text_path():
    segments = extract_text("text/markdown", b"# Heading\n\nBody text.")

    assert len(segments) == 1
    assert "Body text." in segments[0].text


def test_extract_text_plain_skips_blank_content():
    assert extract_text("text/plain", b"   \n\n  ") == []


def test_extract_text_docx_returns_one_segment_per_nonblank_paragraph():
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("")
    document.add_paragraph("Second paragraph.")
    buf = BytesIO()
    document.save(buf)

    segments = extract_text(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document", buf.getvalue()
    )

    assert [s.text for s in segments] == ["First paragraph.", "Second paragraph."]
    assert [s.source_location for s in segments] == [{"paragraph": 0}, {"paragraph": 2}]


def test_extract_text_pdf_returns_one_segment_per_nonblank_page():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = BytesIO()
    writer.write(buf)

    segments = extract_text("application/pdf", buf.getvalue())

    # A blank page has no extractable text, so it contributes no segment.
    assert segments == []


def test_extract_text_rejects_unsupported_mime_type():
    with pytest.raises(UnsupportedMimeTypeError):
        extract_text("application/zip", b"data")
